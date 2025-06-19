import sys
import os
import shutil
import io
import time
import traceback # 导入用于打印详细错误信息的库
import subprocess # 导入用于打开文件夹的库

from PyQt6.QtWidgets import (QApplication, QWidget, QVBoxLayout, QHBoxLayout, 
                             QPushButton, QLineEdit, QListWidget, QLabel,
                             QFileDialog, QMessageBox, QProgressBar, QStyle,
                             QMenu, QInputDialog, QDialog, QTextBrowser) # 新增导入 QDialog, QTextBrowser
from PyQt6.QtCore import Qt
from PyQt6.QtGui import QIcon, QAction, QPixmap # 新增导入 QPixmap

# --- 导入所有处理库 ---
from pypdf import PdfWriter, PdfReader
from PIL import Image
import win32com.client
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_LEADER
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

try:
    pdfmetrics.registerFont(TTFont('SimHei', 'C:/Windows/Fonts/simhei.ttf'))
    pdfmetrics.registerFont(TTFont('SimSun', 'C:/Windows/Fonts/simsun.ttc'))
except Exception as e:
    print(f"字体注册警告: {e}")

# ==================== 新增功能：关于对话框 ====================
class AboutDialog(QDialog):
    """
    一个自定义的对话框，用于显示“关于”信息，包括支持话术、收款码和GitHub链接。
    """
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle('关于 & 支持作者')
        self.setFixedSize(400, 580) # 固定大小以保持界面美观
        self.setStyleSheet(parent.styleSheet()) # 继承主窗口的样式

        layout = QVBoxLayout(self)
        layout.setAlignment(Qt.AlignmentFlag.AlignHCenter)
        layout.setSpacing(15)
        layout.setContentsMargins(20, 20, 20, 20)

        # 顶部话术
        main_text_browser = QTextBrowser()
        main_text_browser.setOpenExternalLinks(True) # 让链接可以被点击
        main_text_browser.setStyleSheet("background-color: transparent; border: none;")
        main_text_browser.setHtml("""
            <p style='text-align:center; font-size:15px;'>
            感谢您使用<b>夏令营申请材料生成器</b>！<br>
            这个小工具旨在简化繁琐的材料准备过程，希望为您节省时间。
            </p>
            <p style='text-align:center; font-size:15px;'>
            如果觉得好用，不妨扫描下方的二维码请作者喝杯奶茶~<br>
            您的支持是我持续更新和维护的最大动力！
            </p>
        """)

        # 收款二维码图片
        qr_label = QLabel()
        # 从本地加载图片，请确保 payment_qr.jpg 和 main.py 在同一目录下
        qr_pixmap = QPixmap('payment_qr.jpg')
        if not qr_pixmap.isNull():
            # 缩放图片并保持比例
            qr_label.setPixmap(qr_pixmap.scaled(250, 250, Qt.AspectRatioMode.KeepAspectRatio, Qt.TransformationMode.SmoothTransformation))
            qr_label.setAlignment(Qt.AlignmentFlag.AlignHCenter)
        else:
            # 如果图片加载失败，显示提示文字
            qr_label.setText("payment_qr.jpg 图片未找到\n请确保它和主程序在同一目录下")
            qr_label.setAlignment(Qt.AlignmentFlag.AlignHCenter)

        # GitHub仓库链接
        github_link_browser = QTextBrowser()
        github_link_browser.setOpenExternalLinks(True)
        github_link_browser.setStyleSheet("background-color: transparent; border: none;")
        github_link_browser.setHtml("""
            <p style='text-align:center; font-size:14px;'>
            本项目的代码已在GitHub开源，欢迎 Star、Fork 或提出 Issues！<br>
            <a style="color: #88C0D0; text-decoration:none;" href='https://github.com/Uruze/Summer-Camp-Material-Generator'>https://github.com/Uruze/Summer-Camp-Material-Generator</a>
            </p>
        """)
        
        # 关闭按钮
        close_button = QPushButton('关闭')
        close_button.clicked.connect(self.accept) # accept() 是 QDialog 的一个槽，用于关闭对话框

        # 将所有控件添加到布局中
        layout.addWidget(main_text_browser)
        layout.addWidget(qr_label)
        layout.addWidget(github_link_browser)
        layout.addWidget(close_button, alignment=Qt.AlignmentFlag.AlignHCenter)

class MainWindow(QWidget):
    def __init__(self):
        super().__init__()
        self.available_file_paths = {}
        self.initUI()

    def initUI(self):
        self.setWindowTitle('夏令营申请材料生成器 by 法国就是培根 (v3.2 - 稳定版)')
        self.resize(900, 700)

        self.setStyleSheet("""
            QWidget {
                background-color: #2E3440;
                color: #D8DEE9;
                font-family: "Microsoft YaHei", "Segoe UI", "Arial", sans-serif;
                font-size: 14px;
            }
            QLabel { color: #E5E9F0; font-size: 15px; }
            QLabel#titleLabel { font-size: 16px; font-weight: bold; color: #88C0D0; padding-bottom: 5px; }
            QLineEdit { background-color: #3B4252; border: 1px solid #4C566A; padding: 8px; border-radius: 5px; color: #ECEFF4; }
            QLineEdit:focus { border: 1px solid #88C0D0; }
            QListWidget { background-color: #3B4252; border: 1px solid #4C566A; border-radius: 5px; padding: 5px; }
            QListWidget::item { padding: 8px; }
            QListWidget::item:hover { background-color: #434C5E; }
            QListWidget::item:selected { background-color: #5E81AC; color: #ECEFF4; border-radius: 3px; }
            QPushButton { background-color: #5E81AC; color: #ECEFF4; border: none; padding: 10px 20px; border-radius: 5px; font-weight: bold; }
            QPushButton:hover { background-color: #81A1C1; }
            QPushButton:pressed { background-color: #4C566A; }
            QProgressBar { border: 1px solid #4C566A; border-radius: 5px; text-align: center; background-color: #3B4252; color: #ECEFF4; }
            QProgressBar::chunk { background-color: #A3BE8C; border-radius: 4px; }
            QMessageBox, QDialog { background-color: #3B4252; }
            QMenu { background-color: #3B4252; border: 1px solid #4C566A; padding: 5px; }
            QMenu::item { padding: 8px 25px; }
            QMenu::item:selected { background-color: #5E81AC; }
            QMenu::separator { height: 1px; background: #4C566A; margin: 5px 0; }
            QTextBrowser { background-color: transparent; border: none; }
        """)

        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(20, 20, 20, 20)
        main_layout.setSpacing(15)

        top_layout = QHBoxLayout()
        school_label = QLabel('申请学校名称:')
        self.school_name_input = QLineEdit()
        self.school_name_input.setPlaceholderText('将替换Word模板中的【目标院校名称】')
        top_layout.addWidget(school_label)
        top_layout.addWidget(self.school_name_input)
        
        core_layout = QHBoxLayout()
        core_layout.setSpacing(15)

        left_layout = QVBoxLayout()
        available_label = QLabel('可用材料库 (可拖拽)')
        available_label.setObjectName("titleLabel")
        self.available_files_list = QListWidget()
        self.available_files_list.setDragEnabled(True)
        self.available_files_list.setSelectionMode(QListWidget.SelectionMode.ExtendedSelection)
        left_layout.addWidget(available_label)
        left_layout.addWidget(self.available_files_list)
        
        right_layout = QVBoxLayout()
        final_label = QLabel('最终材料顺序 (可拖拽排序)')
        final_label.setObjectName("titleLabel")
        self.final_files_list = QListWidget()
        self.final_files_list.setAcceptDrops(True)
        self.final_files_list.setDragEnabled(True)
        self.final_files_list.setDefaultDropAction(Qt.DropAction.MoveAction)
        self.final_files_list.setSelectionMode(QListWidget.SelectionMode.ExtendedSelection)
        right_layout.addWidget(final_label)
        right_layout.addWidget(self.final_files_list)
        
        core_layout.addLayout(left_layout, stretch=1)
        core_layout.addLayout(right_layout, stretch=1)
        
        # --- 修改底部按钮布局 ---
        bottom_layout = QHBoxLayout()
        self.load_button = QPushButton('添加材料文件夹')
        self.about_button = QPushButton('关于 & 支持') # 新增按钮
        self.generate_button = QPushButton('一键生成PDF')
        
        icon_folder = self.style().standardIcon(QStyle.StandardPixmap.SP_DirIcon)
        icon_generate = self.style().standardIcon(QStyle.StandardPixmap.SP_DialogSaveButton)
        icon_about = self.style().standardIcon(QStyle.StandardPixmap.SP_DialogHelpButton)
        self.load_button.setIcon(icon_folder)
        self.generate_button.setIcon(icon_generate)
        self.about_button.setIcon(icon_about) # 为新按钮设置图标

        bottom_layout.addWidget(self.load_button)
        bottom_layout.addStretch() # 使用伸缩项将按钮分开
        bottom_layout.addWidget(self.about_button) # 添加新按钮到布局
        bottom_layout.addWidget(self.generate_button)
        # --- 布局修改结束 ---

        self.progress_bar = QProgressBar()
        self.progress_bar.setTextVisible(True)
        self.progress_bar.setVisible(False)

        main_layout.addLayout(top_layout)
        main_layout.addLayout(core_layout)
        main_layout.addLayout(bottom_layout)
        main_layout.addWidget(self.progress_bar)
        
        # --- 连接按钮信号 ---
        self.load_button.clicked.connect(self.load_materials_folder)
        self.generate_button.clicked.connect(self.generate_final_pdf)
        self.about_button.clicked.connect(self.show_about_dialog) # 连接新按钮的点击信号
        # --- 信号连接结束 ---

        self.available_files_list.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.final_files_list.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.available_files_list.customContextMenuRequested.connect(self.show_context_menu)
        self.final_files_list.customContextMenuRequested.connect(self.show_context_menu)

        self.show()

    def show_about_dialog(self):
        """显示'关于'对话框的槽函数"""
        dialog = AboutDialog(self)
        dialog.exec()

    def show_context_menu(self, position):
        """当在列表上右键单击时，创建并显示菜单"""
        list_widget = self.sender() 
        item = list_widget.itemAt(position)
        if not item:
            return

        menu = QMenu()
        style = self.style()

        action_preview = QAction(style.standardIcon(QStyle.StandardPixmap.SP_FileIcon), "查看预览", self)
        action_open_folder = QAction(style.standardIcon(QStyle.StandardPixmap.SP_DirIcon), "在文件夹中打开", self)
        action_rename = QAction("重命名", self) 

        menu.addAction(action_preview)
        menu.addAction(action_open_folder)
        menu.addSeparator() 
        menu.addAction(action_rename)

        action_preview.triggered.connect(lambda: self.preview_file(item))
        action_open_folder.triggered.connect(lambda: self.open_in_folder(item))
        action_rename.triggered.connect(lambda: self.rename_file(item))

        menu.exec(list_widget.mapToGlobal(position))

    def preview_file(self, item):
        """用系统默认程序打开并预览文件"""
        file_name = item.text()
        file_path = self.available_file_paths.get(file_name)
        if not file_path or not os.path.exists(file_path):
            QMessageBox.warning(self, "错误", f"文件 '{file_name}' 不存在或路径无效。")
            return
        try:
            os.startfile(file_path)
        except Exception as e:
            QMessageBox.critical(self, "预览失败", f"无法打开文件：\n{e}")

    def open_in_folder(self, item):
        """在文件浏览器中打开文件所在的文件夹并选中该文件"""
        file_name = item.text()
        file_path = self.available_file_paths.get(file_name)
        if not file_path or not os.path.exists(file_path):
            QMessageBox.warning(self, "错误", f"文件 '{file_name}' 不存在或路径无效。")
            return
        
        subprocess.run(['explorer', '/select,', os.path.normpath(file_path)])

    def rename_file(self, item):
        """重命名文件，并同步更新程序内的所有引用"""
        old_name_with_ext = item.text()
        old_full_path = self.available_file_paths.get(old_name_with_ext)

        if not old_full_path:
            QMessageBox.critical(self, "错误", "找不到文件的内部记录。")
            return

        old_name_no_ext, ext = os.path.splitext(old_name_with_ext)
        
        new_name_no_ext, ok = QInputDialog.getText(self, '重命名文件', '请输入新的文件名 (不含扩展名):', QLineEdit.EchoMode.Normal, old_name_no_ext)

        if ok and new_name_no_ext and new_name_no_ext != old_name_no_ext:
            invalid_chars = r'<>:"/\|?*'
            if any(char in invalid_chars for char in new_name_no_ext):
                QMessageBox.warning(self, "无效名称", f"文件名不能包含以下任何字符: {invalid_chars}")
                return

            new_name_with_ext = new_name_no_ext + ext
            dir_path = os.path.dirname(old_full_path)
            new_full_path = os.path.join(dir_path, new_name_with_ext)

            if os.path.exists(new_full_path):
                QMessageBox.warning(self, "重命名失败", "同名文件已存在。")
                return

            try:
                os.rename(old_full_path, new_full_path)
            except OSError as e:
                QMessageBox.critical(self, "错误", f"无法在磁盘上重命名文件：\n{e}")
                return
            
            self.available_file_paths.pop(old_name_with_ext)
            self.available_file_paths[new_name_with_ext] = new_full_path
            
            item.setText(new_name_with_ext)
            
            source_list = self.sender()
            other_list = self.final_files_list if source_list == self.available_files_list else self.available_files_list
            
            for i in range(other_list.count()):
                if other_list.item(i).text() == old_name_with_ext:
                    other_list.item(i).setText(new_name_with_ext)
                    break
            
            QMessageBox.information(self, "成功", f"文件已重命名为 '{new_name_with_ext}'")

    def load_materials_folder(self):
        folder_path = QFileDialog.getExistingDirectory(self, '请选择您的材料所在的文件夹')
        if folder_path:
            supported_extensions = ['.pdf', '.docx', '.doc', '.jpg', '.jpeg', '.png']
            for file_name in os.listdir(folder_path):
                if file_name not in self.available_file_paths:
                    if any(file_name.lower().endswith(ext) for ext in supported_extensions):
                        self.available_files_list.addItem(file_name)
                        full_path = os.path.join(folder_path, file_name)
                        self.available_file_paths[file_name] = full_path

    def convert_image_to_pdf(self, img_path, pdf_path):
        try:
            c = canvas.Canvas(pdf_path, pagesize=A4)
            A4_WIDTH, A4_HEIGHT = A4
            
            with Image.open(img_path) as img:
                img_width, img_height = img.size
                max_width = A4_WIDTH - 1 * inch
                max_height = A4_HEIGHT - 1 * inch
                ratio = min(max_width / img_width, max_height / img_height)
                new_width = img_width * ratio
                new_height = img_height * ratio
                x_centered = (A4_WIDTH - new_width) / 2
                y_centered = (A4_HEIGHT - new_height) / 2
                c.drawImage(img_path, x_centered, y_centered, width=new_width, height=new_height, preserveAspectRatio=True)
                
            c.save()
            return True
        except Exception as e:
            raise Exception(f"转换图片失败 {os.path.basename(img_path)}: {e}")

    def convert_word_to_pdf(self, word_path, pdf_path):
        word = None
        doc = None
        try:
            word_path = os.path.abspath(word_path)
            pdf_path = os.path.abspath(pdf_path)
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(word_path)
            doc.SaveAs(pdf_path, FileFormat=17)
            return True
        except Exception as e:
            raise Exception(f"转换Word失败 {os.path.basename(word_path)}: {e}")
        finally:
            if doc: doc.Close(False)
            if word: word.Quit()

    def generate_final_pdf(self):
        school_name = self.school_name_input.text().strip()
        if not school_name:
            QMessageBox.warning(self, '提示', '请输入申请学校的名称！')
            return

        final_items = [self.final_files_list.item(i).text() for i in range(self.final_files_list.count())]
        if not final_items:
            QMessageBox.warning(self, '提示', '请将材料拖拽到右侧的最终顺序列表中！')
            return
        
        template_path = os.path.join(os.path.dirname(sys.argv[0]), "surface.docx")
        if not os.path.exists(template_path):
            QMessageBox.critical(self, '错误', f'未找到模板文件 surface.docx！\n请确保它和 main.py 在同一目录下。')
            return

        save_path, _ = QFileDialog.getSaveFileName(self, '保存文件', f"{school_name}-申请材料.pdf", "PDF Files (*.pdf)")
        if not save_path:
            return

        self.progress_bar.setVisible(True)
        self.progress_bar.setValue(0)
        QApplication.processEvents()
        temp_dir = os.path.join(os.path.dirname(save_path), "temp_conversion")

        try:
            if os.path.exists(temp_dir): shutil.rmtree(temp_dir)
            os.makedirs(temp_dir)
            
            total_steps = len(final_items) + 4
            current_step = 0
            
            converted_pdf_paths = []
            for item_name in final_items:
                current_step += 1
                self.progress_bar.setFormat(f"正在处理: {item_name}")
                self.progress_bar.setValue(int(current_step / total_steps * 100))
                QApplication.processEvents()
                original_path = self.available_file_paths[item_name]
                file_ext = os.path.splitext(item_name)[1].lower()
                temp_pdf_path = os.path.join(temp_dir, f"{current_step}_{os.path.basename(item_name)}.pdf")

                if file_ext == '.pdf':
                    shutil.copy(original_path, temp_pdf_path)
                    converted_pdf_paths.append(temp_pdf_path)
                elif file_ext in ['.jpg', '.jpeg', '.png']:
                    self.convert_image_to_pdf(original_path, temp_pdf_path)
                    converted_pdf_paths.append(temp_pdf_path)
                elif file_ext in ['.doc', '.docx']:
                    self.convert_word_to_pdf(original_path, temp_pdf_path)
                    converted_pdf_paths.append(temp_pdf_path)

            current_step += 1
            self.progress_bar.setFormat("合并内容并统一页面方向...")
            self.progress_bar.setValue(int(current_step / total_steps * 100))
            QApplication.processEvents()
            
            content_merger = PdfWriter()
            toc_entries = []
            toc_page_count = 1
            current_page_in_content = 1
            
            for i, path in enumerate(converted_pdf_paths):
                try:
                    reader = PdfReader(path)
                    num_pages_in_file = len(reader.pages)
                    toc_entries.append({'title': os.path.splitext(final_items[i])[0], 'page': current_page_in_content + toc_page_count})
                    for page in reader.pages:
                        if page.mediabox.width > page.mediabox.height:
                            page.rotate(90)
                        content_merger.add_page(page)
                    current_page_in_content += num_pages_in_file
                except Exception as file_error:
                    raise Exception(f"处理文件 '{os.path.basename(path)}' 时发生错误，文件可能已损坏。") from file_error
            
            content_path = os.path.join(temp_dir, "content.pdf")
            with open(content_path, "wb") as f_content:
                content_merger.write(f_content)
            
            current_step += 1
            self.progress_bar.setFormat("创建精美目录页...")
            self.progress_bar.setValue(int(current_step / total_steps * 100))
            QApplication.processEvents()
            doc = Document(template_path)
            for p in doc.paragraphs:
                if '【目标院校名称】' in p.text:
                    p.text = ""
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(school_name)
                    run.font.name = '黑体'
                    run.font.size = Pt(26)
                    break
            for p in doc.paragraphs:
                if '【目录】' in p.text:
                    p.text = ""
                    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.0), WD_TAB_LEADER.DOTS, WD_ALIGN_PARAGRAPH.RIGHT)
                    for entry in toc_entries:
                        run = p.add_run(f"{entry['title']}\t{entry['page']}")
                        font = run.font; font.name = '宋体'; font.size = Pt(12)
                        p.add_run().add_break()
                    break
            temp_toc_docx_path = os.path.join(temp_dir, "toc_temp.docx")
            doc.save(temp_toc_docx_path)
            temp_toc_pdf_path = os.path.join(temp_dir, "toc_temp.pdf")
            self.convert_word_to_pdf(temp_toc_docx_path, temp_toc_pdf_path)

            current_step += 1
            self.progress_bar.setFormat("最终合并并添加页码...")
            self.progress_bar.setValue(int(current_step / total_steps * 100))
            QApplication.processEvents()
            
            final_merger = PdfWriter()
            final_merger.append(temp_toc_pdf_path)
            final_merger.append(content_path)

            for i, page in enumerate(final_merger.pages):
                packet = io.BytesIO()
                page_width = page.mediabox.width
                page_height = page.mediabox.height
                c = canvas.Canvas(packet, pagesize=(page_width, page_height))
                c.setFont('Helvetica', 9)
                page_number_text = str(i + 1)
                c.drawCentredString(float(page_width) / 2, 0.5 * inch, page_number_text)
                c.save()
                packet.seek(0)
                watermark_pdf = PdfReader(packet)
                page.merge_page(watermark_pdf.pages[0])

            with open(save_path, "wb") as f: final_merger.write(f)
            self.progress_bar.setValue(100)
            self.progress_bar.setFormat("完成!")
            QMessageBox.information(self, '成功', f'文件已成功生成！\n保存在: {save_path}')

        except Exception as e:
            traceback.print_exc() 
            QMessageBox.critical(self, '发生错误', f"生成过程中出现问题，操作已中断。\n\n错误信息:\n{e}")
        
        finally:
            if os.path.exists(temp_dir):
                shutil.rmtree(temp_dir)
            self.progress_bar.setVisible(False)

if __name__ == '__main__':
    # 确保在高DPI屏幕上显示正常
    if hasattr(Qt, 'AA_EnableHighDpiScaling'):
        QApplication.setAttribute(Qt.ApplicationAttribute.AA_EnableHighDpiScaling, True)
    if hasattr(Qt, 'AA_UseHighDpiPixmaps'):
        QApplication.setAttribute(Qt.ApplicationAttribute.AA_UseHighDpiPixmaps, True)
    
    app = QApplication(sys.argv)
    main_window = MainWindow()
    sys.exit(app.exec())